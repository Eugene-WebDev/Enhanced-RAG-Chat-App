import os
import io
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd
import markdown

def extract_text_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_txt(file):
    return file.getvalue().decode("utf-8")

def extract_text_docx(file):
    # Save file temporarily in memory
    file_bytes = io.BytesIO(file.getvalue())
    doc = DocxDocument(file_bytes)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_csv(file):
    # Using pandas to read CSV and then join all cell values as text
    df = pd.read_csv(file)
    return df.to_csv(index=False)

def extract_text_md(file):
    # Markdown files: decode and optionally convert to plain text
    md_text = file.getvalue().decode("utf-8")
    # Optionally, convert markdown to HTML or leave as plain text
    html = markdown.markdown(md_text)
    return html

def extract_document_text(file):
    filename = file.name.lower()
    if filename.endswith(".pdf"):
        from PyPDF2 import PdfReader
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    elif filename.endswith(".txt"):
        return file.getvalue().decode("utf-8")
    elif filename.endswith(".docx"):
        from docx import Document as DocxDocument
        import io
        file_bytes = io.BytesIO(file.getvalue())
        doc = DocxDocument(file_bytes)
        return "\n".join([para.text for para in doc.paragraphs])
    # Add other file types as needed
    else:
        raise ValueError(f"Unsupported file type: {file.name}")